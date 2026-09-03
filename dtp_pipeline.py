#!/usr/bin/env python3
"""
dtp_pipeline.py - Universal DOCX DTP / standardization pipeline.

Takes a raw, unedited Word document + a JSON "design system" file and produces
a professionally formatted DOCX: consistent fonts/colors, styled headings,
styled tables (vocabulary tables, dialogue tables, self-assessment grids,
word-search grids, blank-answer lines), a running header/footer, and
correctly-sized re-inserted images -- all driven by the JSON config, not by
anything hard-coded to one particular document.

USAGE
-----
    python dtp_pipeline.py INPUT.docx
    python dtp_pipeline.py INPUT.docx -s style.json -o OUTPUT.docx
    python dtp_pipeline.py INPUT.docx --pdf          # also render a PDF via LibreOffice

If -s/--style is omitted, a sensible built-in default design system is used
(A4, Google-blue palette, Noto Sans / Arial). Any JSON you do pass only needs
to override the keys you care about -- everything else falls back to the
default.

HOW CLASSIFICATION WORKS (no manual tagging needed)
----------------------------------------------------
The script walks the source document top-to-bottom (paragraphs AND tables,
in their original order) and classifies each block using structural signals
that are already present in almost any worksheet/workbook-style document:

  * bold + short text                       -> heading / sub-heading
  * "Unit 1:", "REVISION", "Skit:", etc.     -> title
  * a curated list of common worksheet
    section names ("Objective:", "New
    words...", "Written Exercises", ...)     -> section heading
  * "Name: some sentence" pattern            -> dialogue line / dialogue table
  * a table whose first row is entirely
    bold and whose second row is NOT         -> vocabulary/definition table
    (a strong, language-independent signal
    of "this row is a header")
  * 10x10 grid of single characters          -> word-search grid
  * table with completely empty cells        -> ruled answer-lines

None of this is specific to "Grade 7 Spoken English" -- it is pattern-based,
so pointing the script at a different raw DOCX of similar structure (a
worksheet / workbook with units, vocab tables, dialogues, exercises) will work
without touching the code. The SECTION_KEYWORDS list below is the one part
you may want to extend if your documents use different section names.

REQUIREMENTS
------------
    pip install python-docx
Optional, only for --pdf:
    LibreOffice ('soffice') installed and on PATH.
"""

import argparse
import copy
import os
import re
import subprocess
import sys
import tempfile

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    import json
except ImportError:  # pragma: no cover
    raise

try:
    from PIL import Image, ImageChops
    HAVE_PIL = True
except ImportError:  # pragma: no cover
    HAVE_PIL = False


# Folder for bundled cover/footer artwork (logos, illustration) shipped
# alongside this script. Keep the "assets" folder next to dtp_pipeline.py.
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(SCRIPT_DIR, "assets")


# ============================================================================
# 0. >>> EDIT THESE, THEN JUST RUN:  python3 dtp_pipeline.py <<<
# ============================================================================
# This is the only section you normally need to touch. Point INPUT_DOCX at
# your raw/unedited Word file and run the script with no arguments -- it will
# use the settings below automatically.
#
#   INPUT_DOCX   - path to the raw .docx you want formatted (required)
#   OUTPUT_DOCX  - path to write the formatted .docx to.
#                  Leave as None to auto-name it "<input>_formatted.docx"
#                  next to the input file.
#   STYLE_JSON   - path to a JSON design-system file (colors/fonts/tables/
#                  footer text). Leave as None to use the built-in defaults
#                  further down in this file.
#   MAKE_PDF     - True to also render a PDF via LibreOffice ('soffice')
#                  after formatting, False to skip it.
#
# (All of these can still be overridden from the command line if you want,
#  e.g. `python3 dtp_pipeline.py other_file.docx -o custom_output.docx` --
#  but for normal day-to-day use you can just edit the values below and run
#  the script with no arguments at all.)
# ----------------------------------------------------------------------------

INPUT_DOCX = "input.docx"
OUTPUT_DOCX = None
STYLE_JSON = None
MAKE_PDF = False

# ----------------------------------------------------------------------------
# OPTIONAL FRONT MATTER  (cover page / About / How-to-use / Contents page)
# ----------------------------------------------------------------------------
# All of this is OFF by default (COVER["enabled"] = False) so a bare run just
# formats the body content, exactly like before. Turn it on and fill in the
# fields below once for a given book series, and every future run of the
# script will regenerate the same cover/front-matter automatically.
#
# Logos/illustrations are real image files the raw DOCX does not necessarily
# contain (a cover is a one-time design asset, not something that can be
# invented from arbitrary body content) -- point the paths below at your own
# PNG/JPG files. Any path left empty ("") or pointing at a missing file is
# simply skipped, so it's safe to leave items blank.

COVER = {
    "enabled": True,
    "logos": [
        os.path.join(ASSETS_DIR, "logo_sikshana.png"),
        os.path.join(ASSETS_DIR, "logo_karnataka_emblem.png"),
        os.path.join(ASSETS_DIR, "logo_qualcomm.png"),
    ],
    "illustration": os.path.join(ASSETS_DIR, "cover_illustration.png"),
    "title": "Workbook",         # big title text -- edit per book
    "subtitle": "Class VII",     # e.g. class/level line -- edit per book
    "fields": ["Name", "Std", "School"],   # blank ruled fields printed below the title
}

# Static introductory sections (title + paragraphs), inserted after the cover
# and before the auto-generated Contents page. Leave the list empty to skip.
FRONT_MATTER = [
    # {"heading": "About this Book", "paragraphs": ["...", "..."]},
    # {"heading": "How to Use this Book", "paragraphs": ["...", "..."]},
]

# Auto-builds a Contents page from every detected Unit/title heading, using a
# native Word TOC field (Word: right-click it -> "Update Field" once to fill
# in page numbers; this is standard Word behaviour for any generated TOC).
ADD_CONTENTS_PAGE = True

# ============================================================================
# 1. DEFAULT DESIGN SYSTEM  (used verbatim, or overridden by STYLE_JSON above)
# ============================================================================

DEFAULT_CONFIG = {
    "document_setup": {
        "width_mm": 210, "height_mm": 297,
        "margins_mm": {"top": 20, "bottom": 20, "left": 15, "right": 15},
    },
    "color_palette": {
        # main heading / title color
        "primary": "#D1232A",
        # section-heading text AND table-header background ("table blue")
        "secondary": "#2596BE",
        "dark_neutral": "#202124",
        # light purple -- used as the "heading highlight" / objective box bg
        "light_neutral": "#DCDDDF",
        "accent_bg": "#DCDDDF",
        # green table shading + orange borders, as requested
        "table_alt_green": "#D8E9DE",
        "border_color": "#ECB098",
        "border_color_alt": "#E9A791",
        "text_muted": "#5F6368",
    },
    "typography": {
        # ascii/Latin font. NOTE: Bookman Old Style has no Kannada/Indic
        # glyphs, so Kannada/Devanagari text automatically uses "indic"
        # below instead -- this keeps both scripts legible.
        "font_family": {"primary": "Bookman Old Style", "fallback": "Georgia",
                        "indic": "Nirmala UI"},
        "styles": {
            "document_header": {"font_size_pt": 10, "font_weight": "bold",
                                 "alignment": "right",
                                 "format": "Date: ____ / ____ / ____"},
            "unit_title": {"font_size_pt": 20, "font_weight": "bold",
                           "color": "#D1232A", "margin_bottom_pt": 15},
            "objective_box": {"font_size_pt": 10.5, "font_weight": "bold",
                               "background_color": "#DCDDDF",
                               "border_left_width_pt": 4,
                               "border_left_color": "#E9A791",
                               "padding_pt": 8},
            "section_heading": {"font_size_pt": 13, "font_weight": "bold",
                                 "color": "#2596BE",
                                 "border_bottom_width_pt": 1,
                                 "border_bottom_color": "#ECB098",
                                 "margin_top_pt": 18, "margin_bottom_pt": 8},
            "body_text": {"font_size_pt": 10.5, "line_height_pt": 14,
                          "color": "#202124"},
            "kannada_translation": {"font_size_pt": 9.5, "line_height_pt": 13,
                                     "color": "#5F6368"},
        },
    },
    "components": {
        "tables": {
            "vocabulary_table": {"column_width_ratios": [0.45, 0.55],
                                  "header_bg": "#2596BE",
                                  "header_text_color": "#FFFFFF",
                                  "alternate_row_bg": "#D8E9DE",
                                  "border_width_pt": 0.5, "cell_padding_pt": 8},
            "grammar_table": {"column_width_ratios": [0.3, 0.3, 0.4],
                               "header_bg": "#2596BE",
                               "header_text_color": "#FFFFFF",
                               "border_width_pt": 0.5, "cell_padding_pt": 8},
        },
    },
    "footer": {
        "left_text": "Document", "right_text": "Page {page_num}",
        "font_size_pt": 9, "color": "#555555",
        # NEW: put a logo in the center of the footer instead of left text,
        # and push the page number into a page corner. Leave "logo" empty
        # ("") to fall back to the old left-text-plus-page-number layout.
        "logo": os.path.join(ASSETS_DIR, "logo_sikshana.png"),
        "logo_height_cm": 0.7,
        "page_number_corner": "right",   # "right" or "left"
    },
}


def deep_merge(base, override):
    """Recursively merge `override` on top of `base`; returns a new dict."""
    result = copy.deepcopy(base)
    for k, v in (override or {}).items():
        if isinstance(v, dict) and isinstance(result.get(k), dict):
            result[k] = deep_merge(result[k], v)
        else:
            result[k] = v
    return result


# ============================================================================
# 2. LOW-LEVEL DOCX/XML HELPERS
# ============================================================================

def hexcolor(h):
    h = (h or "000000").lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_font(run, name, size=10.5, bold=None, italic=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    pPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        node = OxmlElement(f'w:{tag}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def set_table_borders(table, color='BDC1C6', sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color.lstrip('#'))
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, ratios, total_cm=18.0):
    total_dxa = int(Cm(total_cm))
    table.autofit = False
    widths = [int(total_dxa * r) for r in ratios]
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gc, w in zip(tblGrid.findall(qn('w:gridCol')), widths):
            gc.set(qn('w:w'), str(w))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')


def para_spacing(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)


def add_bottom_border(paragraph, color='DADCE0', sz=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color.lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_left_border(paragraph, color, sz=24):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color.lstrip('#'))
    pBdr.append(left)
    pPr.append(pBdr)


def add_page_number_field(paragraph, size=9, color=None):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    set_font(run, name=run.font.name or 'Arial', size=size, color=color)


def set_row_cant_split(row):
    """Ask Word/LibreOffice not to break this table row across a page boundary."""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:cantSplit')
    trPr.append(el)


def insert_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return p


def add_toc_field(paragraph):
    """Insert a native Word Table-of-Contents field (levels 1-1)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-1" \\h \\z \\u '
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Right-click here and choose \u201cUpdate Field\u201d to build the table of contents."
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


# ============================================================================
# 3. CLASSIFICATION HEURISTICS  (pattern based, not content specific)
# ============================================================================

TITLE_RE = re.compile(
    r'^(UNIT|Unit|CHAPTER|Chapter|LESSON|Lesson)\s*\d+|^REVISION$|^Skit:|'
    r'^Word Search!$|^The End!?$', re.IGNORECASE)
SCENE_RE = re.compile(r'^Scene\s+\d+', re.IGNORECASE)
LETTERED_RE = re.compile(r'^[A-E]\.\s+\S')
DIALOGUE_RE = re.compile(r"^(\d+\.\s*)?[A-Za-z][A-Za-z .\u2019'()]{0,30}:\s+\S")
KANNADA_RE = re.compile(r'[\u0C80-\u0CFF]')
DEVANAGARI_RE = re.compile(r'[\u0900-\u097F]')
NON_LATIN_RE = re.compile(r'[\u0900-\u0DFF]')  # broad Indic-script range

OBJECTIVE_INLINE_RE = re.compile(r'^(Objective|Learning Objective)\s*:\s*\S', re.IGNORECASE)

# Common worksheet/workbook section names. Extend this list for other
# document families -- everything else in the pipeline is pattern based.
SECTION_KEYWORDS = [
    'objective:', 'learning objective', 'new words', 'sentences to learn',
    'phrases to learn', 'practice dialogue', 'dialogue practice',
    'natural conversation practice', 'role-play activit', 'audio activity',
    'listen carefully', 'listen and write true or false', 'written exercises',
    'fill in the blanks', 'arrange the words', 'speak at home challenge',
    'speaking at home challenge', 'self-assessment', 'translate the words',
    'translate the sentences', 'translate the kannada', 'translate into english',
    'instructions for the facilitator', 'characters', 'practice tips',
    'make your simple study plan', 'now say:', 'then tell your partner:',
    'word bank:', 'use the words in the box.', 'choose the best reply',
    'who should say it', 'which question can you ask', 'match the',
    'put the words in order', 'write easy or difficult', 'find and circle',
    'hint:', 'conversation practice', 'role play', 'self assessment',
    'speaking challenge', 'natural conversation',
]


# Section names whose *following content* (until the next heading/table)
# should be rendered as a shaded callout box instead of plain paragraphs --
# covers "Objective", "Speaking Challenge", "Word bank", etc.
CALLOUT_TRIGGERS = {
    'objective:', 'learning objective', 'speak at home challenge',
    'speaking at home challenge', 'speaking challenge', 'word bank:',
    'use the words in the box.',
}

# Small Unicode glyphs used as lightweight "icons" next to recognized
# section headings (no external image assets required). Purely cosmetic --
# safe to edit/clear per keyword.
HEADING_ICONS = {
    'practice dialogue': '\U0001F5E8 ', 'dialogue practice': '\U0001F5E8 ',
    'audio activity': '\U0001F50A ', 'listen carefully': '\U0001F3A7 ',
    'listen and write true or false': '\U0001F3A7 ',
    'role-play activit': '\U0001F3AD ', 'role play': '\U0001F3AD ',
    'self-assessment': '\u2705 ', 'self assessment': '\u2705 ',
    'word bank:': '\U0001F9E9 ',
    'speak at home challenge': '\U0001F4E2 ', 'speaking at home challenge': '\U0001F4E2 ',
    'speaking challenge': '\U0001F4E2 ',
    'natural conversation practice': '\U0001F4AC ', 'conversation practice': '\U0001F4AC ',
}


def strip_lead_numbering(t):
    return re.sub(r'^(\d{1,2}\.\s*|[A-E]\.\s*)', '', t).strip()


def classify_para(text, src_bold=False):
    t = text.strip()
    if not t:
        return 'skip'
    if TITLE_RE.match(t):
        return 'title'
    if OBJECTIVE_INLINE_RE.match(t):
        return 'objective_inline'
    if SCENE_RE.match(t):
        return 'heading'
    core = strip_lead_numbering(t).lower()
    for kw in SECTION_KEYWORDS:
        if core.startswith(kw):
            return 'heading'
    if NON_LATIN_RE.match(t):
        return 'body'
    if t.count('\u2013') + t.count('-') >= 2 and len(t.split()) <= 12:
        return 'body'
    if DIALOGUE_RE.match(t):
        return 'dialogue_para'
    if LETTERED_RE.match(t):
        return 'subheading'
    words = t.split()
    has_punct_marks = ('/' in t) or ('_' in t) or t[0:1].islower()
    if (src_bold and len(words) <= 9 and not t.endswith('?')
            and not re.search(r'_{3,}', t) and not has_punct_marks):
        return 'subheading'
    return 'body'


def classify_table(table):
    nrows, ncols = len(table.rows), len(table.columns)
    if nrows == 0 or ncols == 0:
        return 'generic'

    all_empty = all(c.text.strip() == '' for r in table.rows for c in r.cells)
    if all_empty:
        return 'blank_lines'

    if ncols == 10 and nrows == 10:
        return 'wordgrid'

    def is_dialogue_cell(t):
        if ':' not in t:
            return False
        name = t.split(':', 1)[0].strip()
        return 0 < len(name.split()) <= 4 and len(t.split()) > len(name.split())

    dialogue_hits = sum(1 for r in table.rows if is_dialogue_cell(r.cells[0].text.strip()))
    if ncols == 2 and dialogue_hits >= max(1, nrows // 2):
        return 'dialogue'

    def cell_is_bold(cell):
        runs = [r for p in cell.paragraphs for r in p.runs if r.text.strip()]
        return bool(runs) and all(r.bold for r in runs)

    row0 = table.rows[0].cells
    header_bold = (all(cell_is_bold(c) for c in row0 if c.text.strip())
                   and any(c.text.strip() for c in row0))
    row1_bold = False
    if nrows > 1:
        row1 = table.rows[1].cells
        row1_bold = all(cell_is_bold(c) for c in row1 if c.text.strip())
    if header_bold and not row1_bold:
        return 'vocab'
    return 'generic'


# ============================================================================
# 4. IMAGE EXTRACTION + AUTO-CROP
# ============================================================================

def get_image_blob(src_doc, p_element):
    xml = p_element.xml
    m = re.search(r'r:embed="(rId\d+)"', xml)
    if not m:
        return None, None, None
    rid = m.group(1)
    wm = re.search(r'<wp:extent cx="(\d+)" cy="(\d+)"', xml)
    w_emu, h_emu = (int(wm.group(1)), int(wm.group(2))) if wm else (None, None)
    try:
        part = src_doc.part.related_parts[rid]
        return part.blob, w_emu, h_emu
    except Exception:
        return None, w_emu, h_emu


def _autocrop_bbox(im, pad_px=6):
    """Return a bounding box (l, t, r, b) that trims uniform-background
    whitespace/transparent margin from a PIL image, plus a small padding."""
    rgba = im.convert("RGBA")
    alpha = rgba.split()[-1]
    bbox = None
    if alpha.getextrema() != (255, 255):
        # real transparency present -> crop to the non-transparent pixels
        bbox = alpha.getbbox()
    if bbox is None:
        # no usable alpha channel -> crop against the image's own corner
        # (background) color, e.g. a plain white/JPEG photo. A small
        # threshold absorbs JPEG noise so faint compression artifacts in
        # the "blank" margin don't fool the bounding box into thinking the
        # whole image is content.
        bg_color = rgba.getpixel((0, 0))
        bg = Image.new("RGBA", rgba.size, bg_color)
        diff = ImageChops.difference(rgba.convert("RGB"), bg.convert("RGB"))
        diff = diff.convert("L").point(lambda px: 255 if px > 18 else 0)
        bbox = diff.getbbox()
    if not bbox:
        return None
    l, t, r, b = bbox
    l = max(0, l - pad_px)
    t = max(0, t - pad_px)
    r = min(im.width, r + pad_px)
    b = min(im.height, b + pad_px)
    if r <= l or b <= t:
        return None
    return (l, t, r, b)


def autocrop_image_file(src_path, dst_path, pad_px=6):
    """Trim whitespace/transparent margin from src_path, save to dst_path.
    Returns (dst_path, cropped_w_px, cropped_h_px, orig_w_px, orig_h_px).
    If PIL isn't available or cropping fails for any reason, the original
    file/size is returned unchanged so the pipeline still runs."""
    if not HAVE_PIL:
        return src_path, None, None, None, None
    try:
        im = Image.open(src_path)
        orig_w, orig_h = im.size
        bbox = _autocrop_bbox(im, pad_px)
        if bbox:
            im.crop(bbox).save(dst_path)
            w, h = im.crop(bbox).size
        else:
            im.save(dst_path)
            w, h = orig_w, orig_h
        return dst_path, w, h, orig_w, orig_h
    except Exception:
        return src_path, None, None, None, None


def cropped_extent(w_emu, h_emu, orig_w_px, orig_h_px, cropped_w_px, cropped_h_px):
    """Rescale an original EMU width/height to match a cropped image, keeping
    the same physical pixels-per-EMU ("print scale") as the original."""
    if not (w_emu and h_emu and orig_w_px and orig_h_px and cropped_w_px and cropped_h_px):
        return w_emu, h_emu
    return (int(w_emu * cropped_w_px / orig_w_px),
            int(h_emu * cropped_h_px / orig_h_px))


# ============================================================================
# 5. MAIN BUILD
# ============================================================================

def build(input_path, output_path, cfg, tmp_dir,
          cover=None, front_matter=None, add_contents=False):
    cover = cover or {"enabled": False}
    front_matter = front_matter or []
    PAL = cfg['color_palette']
    TYP = cfg['typography']
    STY = TYP['styles']
    FONT = TYP['font_family']['primary']
    FONT_INDIC = TYP['font_family'].get('indic', 'Nirmala UI')
    TBL = cfg['components']['tables']
    FOOT = cfg['footer']
    PAGE = cfg['document_setup']

    C_PRIMARY = hexcolor(PAL['primary'])
    C_SECONDARY = hexcolor(PAL['secondary'])
    C_DARK = hexcolor(PAL['dark_neutral'])
    C_MUTED = hexcolor(PAL['text_muted'])

    def font_for(text):
        """Pick the Latin or Indic font depending on the run's script."""
        return FONT_INDIC if NON_LATIN_RE.search(text or '') else FONT

    vocab_cfg = TBL.get('vocabulary_table', DEFAULT_CONFIG['components']['tables']['vocabulary_table'])
    grammar_cfg = TBL.get('grammar_table', DEFAULT_CONFIG['components']['tables']['grammar_table'])
    border_color = PAL.get('border_color', '#BDC1C6')
    border_sz = max(2, int(vocab_cfg.get('border_width_pt', 0.5) * 8))  # pt -> eighths-of-a-point

    src_doc = Document(input_path)
    out = Document()

    # ---- page setup ----
    sec = out.sections[0]
    sec.page_width = Mm(PAGE.get('width_mm', 210))
    sec.page_height = Mm(PAGE.get('height_mm', 297))
    margins = PAGE.get('margins_mm', {})
    sec.top_margin = Mm(margins.get('top', 20))
    sec.bottom_margin = Mm(margins.get('bottom', 20))
    sec.left_margin = Mm(margins.get('left', 15))
    sec.right_margin = Mm(margins.get('right', 15))
    sec.header_distance = Mm(10)
    sec.footer_distance = Mm(10)
    if cover.get('enabled'):
        # keep the running "Date:" header off the cover page itself
        sec.different_first_page_header_footer = True

    # ---- base style ----
    normal = out.styles['Normal']
    normal.font.size = Pt(STY.get('body_text', {}).get('font_size_pt', 10.5))
    normal.font.color.rgb = C_DARK
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)

    # ---- header ----
    hdr_cfg = STY.get('document_header', {})
    hp = sec.header.paragraphs[0]
    align = hdr_cfg.get('alignment', 'right')
    hp.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.RIGHT)
    hr = hp.add_run(hdr_cfg.get('format', ''))
    set_font(hr, FONT, size=hdr_cfg.get('font_size_pt', 10),
             bold=(hdr_cfg.get('font_weight') == 'bold'), color=C_DARK)
    para_spacing(hp, 0, 0)

    # ---- footer ----
    foot_color = hexcolor(FOOT.get('color', '#555555'))
    foot_size = FOOT.get('font_size_pt', 9)
    logo_path = FOOT.get('logo', '')
    page_corner = FOOT.get('page_number_corner', 'right')
    page_align = WD_ALIGN_PARAGRAPH.RIGHT if page_corner == 'left' else WD_ALIGN_PARAGRAPH.LEFT
    # (the corner column sits on the *opposite* side of the table cell it's
    # in, hence right-aligned text inside the "left corner" cell and vice
    # versa -- see the 3-column footer table below)

    if logo_path and os.path.isfile(logo_path):
        # Footer laid out as a single borderless 3-column row so the logo is
        # centered on the page while the page number sits in a corner, both
        # on the same line.
        avail_cm = PAGE.get('width_mm', 210) / 10 - margins.get('left', 15) / 10 - margins.get('right', 15) / 10
        ftbl = sec.footer.add_table(rows=1, cols=3, width=Cm(avail_cm))
        ftbl.autofit = False
        set_col_widths(ftbl, [0.3, 0.4, 0.3], total_cm=avail_cm)
        for row in ftbl.rows:
            for cell in row.cells:
                set_cell_margins(cell, 0, 0, 0, 0)
        left_cell, mid_cell, right_cell = ftbl.rows[0].cells
        corner_cell = left_cell if page_corner == 'left' else right_cell
        other_cell = right_cell if page_corner == 'left' else left_cell

        mid_p = mid_cell.paragraphs[0]
        mid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            cropped = os.path.join(tmp_dir, 'footer_logo.png')
            cfn, cw, ch, ow, oh = autocrop_image_file(logo_path, cropped)
            mid_p.add_run().add_picture(cfn, height=Cm(FOOT.get('logo_height_cm', 0.8)))
        except Exception:
            pass

        corner_p = corner_cell.paragraphs[0]
        corner_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if page_corner == 'right' else WD_ALIGN_PARAGRAPH.LEFT
        add_page_number_field(corner_p, size=foot_size, color=foot_color)

        other_cell.paragraphs[0].text = ''

        # remove the blank paragraph python-docx leaves under the table
        extra_p = sec.footer.paragraphs[-1]
        if extra_p._p.getparent() is not None and not extra_p.text:
            para_spacing(extra_p, 0, 0)
    else:
        # fallback: classic "left text ... tab ... Page N" single-line footer
        fp = sec.footer.paragraphs[0]
        fp.text = ''
        fp.paragraph_format.tab_stops.add_tab_stop(
            Mm(PAGE.get('width_mm', 210) - margins.get('left', 15) - margins.get('right', 15)), alignment=2)
        left_text = FOOT.get('left_text', '')
        right_text = FOOT.get('right_text', 'Page {page_num}')
        r1 = fp.add_run(left_text)
        set_font(r1, FONT, size=foot_size, color=foot_color)
        r1.add_tab()
        before, _, after = right_text.partition('{page_num}')
        if before:
            r = fp.add_run(before)
            set_font(r, FONT, size=foot_size, color=foot_color)
        add_page_number_field(fp, size=foot_size, color=foot_color)
        if after:
            r = fp.add_run(after)
            set_font(r, FONT, size=foot_size, color=foot_color)
        para_spacing(fp, 0, 0)

    # ---- paragraph style renderers ----
    title_seen = {'n': 0}

    def style_title(p, text, page_break=True):
        tcfg = STY.get('unit_title', {})
        if page_break and title_seen['n'] > 0:
            p.paragraph_format.page_break_before = True
        title_seen['n'] += 1
        try:
            p.style = out.styles['Heading 1']  # lets Word's TOC field find this
        except KeyError:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, FONT, size=tcfg.get('font_size_pt', 20),
                  bold=(tcfg.get('font_weight', 'bold') == 'bold'),
                  color=hexcolor(tcfg.get('color', PAL['primary'])))
        para_spacing(p, 6, tcfg.get('margin_bottom_pt', 15),
                     line=tcfg.get('line_height_pt'))
        p.paragraph_format.keep_with_next = True
        add_bottom_border(p, color=tcfg.get('color', PAL['primary']), sz=14)

    def style_heading(p, text):
        hcfg = STY.get('section_heading', {})
        core_lower = strip_lead_numbering(text).lower()
        icon = ''
        for kw, glyph in HEADING_ICONS.items():
            if core_lower.startswith(kw):
                icon = glyph
                break
        r = p.add_run(icon + text)
        set_font(r, FONT, size=hcfg.get('font_size_pt', 13),
                  bold=(hcfg.get('font_weight', 'bold') == 'bold'),
                  color=hexcolor(hcfg.get('color', PAL['secondary'])))
        para_spacing(p, hcfg.get('margin_top_pt', 18), hcfg.get('margin_bottom_pt', 8))
        add_bottom_border(p, color=hcfg.get('border_bottom_color', '#DADCE0'),
                           sz=max(2, int(hcfg.get('border_bottom_width_pt', 1) * 8)))
        p.paragraph_format.keep_with_next = True

    def style_subheading(p, text):
        r = p.add_run(text)
        set_font(r, FONT, size=11.5, bold=True, color=C_SECONDARY)
        para_spacing(p, 10, 4)

    def style_objective(p, label, rest):
        ocfg = STY.get('objective_box', {})
        r1 = p.add_run((label + ' ') if label else '')
        set_font(r1, FONT, size=ocfg.get('font_size_pt', 10.5), bold=True, color=C_DARK)
        r2 = p.add_run(rest)
        set_font(r2, FONT, size=ocfg.get('font_size_pt', 10.5), bold=False, color=C_DARK)
        shade_paragraph(p, ocfg.get('background_color', PAL['accent_bg']))
        add_left_border(p, ocfg.get('border_left_color', PAL['primary']),
                         sz=max(2, int(ocfg.get('border_left_width_pt', 4) * 8)))
        p.paragraph_format.left_indent = Pt(10)
        para_spacing(p, 6, 12, line=STY.get('body_text', {}).get('line_height_pt'))

    def style_body(p, text, bold=False):
        bcfg = STY.get('body_text', {})
        r = p.add_run(text)
        set_font(r, font_for(text), size=bcfg.get('font_size_pt', 10.5), bold=bold,
                  color=hexcolor(bcfg.get('color', PAL['dark_neutral'])))
        para_spacing(p, 0, 6, line=bcfg.get('line_height_pt'))

    DIALOGUE_SPLIT_RE = re.compile(r'^((?:\d+\.\s*)?[A-Za-z][A-Za-z .\u2019\'()]{0,30}:)(\s+)(.*)$')

    def style_dialogue_para(p, text):
        bcfg = STY.get('body_text', {})
        m = DIALOGUE_SPLIT_RE.match(text)
        if m:
            r1 = p.add_run(m.group(1) + ' ')
            set_font(r1, font_for(m.group(1)), size=bcfg.get('font_size_pt', 10.5), bold=True, color=C_PRIMARY)
            r2 = p.add_run(m.group(3))
            set_font(r2, font_for(m.group(3)), size=bcfg.get('font_size_pt', 10.5), color=C_DARK)
        else:
            r = p.add_run(text)
            set_font(r, font_for(text), size=bcfg.get('font_size_pt', 10.5), color=C_DARK)
        para_spacing(p, 0, 5, line=bcfg.get('line_height_pt'))

    def insert_callout_box(items):
        """items: list of (text, bold, is_dialogue) tuples -> one shaded box."""
        if not items:
            return
        ocfg = STY.get('objective_box', {})
        tbl = out.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(tbl, [1.0], total_cm=18.0)
        cell = tbl.rows[0].cells[0]
        shade_cell(cell, ocfg.get('background_color', PAL['accent_bg']))
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'right', 'bottom'):
            el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'), 'nil'); borders.append(el)
        left = OxmlElement('w:left'); left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(max(2, int(ocfg.get('border_left_width_pt', 4) * 8))))
        left.set(qn('w:color'), ocfg.get('border_left_color', PAL['primary']).lstrip('#'))
        borders.append(left)
        tcPr.append(borders)
        set_cell_margins(cell, 120, 120, 200, 160)
        for i, (text, bold, is_dialogue) in enumerate(items):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            if is_dialogue:
                style_dialogue_para(p, text)
            else:
                style_body(p, text, bold=bold)
        spacer = out.add_paragraph()
        para_spacing(spacer, 0, 6)

    def build_cover_page():
        if not cover.get('enabled'):
            return
        logos = [pth for pth in cover.get('logos', []) if pth and os.path.isfile(pth)]
        if logos:
            content_w_cm = PAGE.get('width_mm', 210) / 10 - margins.get('left', 15) / 10 - margins.get('right', 15) / 10
            max_col_w_cm = (content_w_cm / len(logos)) - 0.4  # small gutter per column
            max_h_cm = 2.2
            lt = out.add_table(rows=1, cols=len(logos))
            lt.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, pth in enumerate(logos):
                cell = lt.rows[0].cells[i]
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    cropped = os.path.join(tmp_dir, f'cover_logo_{i}.png')
                    cfn, cw, ch, ow, oh = autocrop_image_file(pth, cropped)
                    # Fit within both a max height AND the column width --
                    # whichever constraint is tighter for this logo's shape.
                    h_cm = max_h_cm
                    if cw and ch:
                        w_at_max_h = max_h_cm * (cw / ch)
                        if w_at_max_h > max_col_w_cm:
                            h_cm = max_col_w_cm * (ch / cw)
                    cp.add_run().add_picture(cfn, height=Cm(h_cm))
                except Exception:
                    pass
            spacer = out.add_paragraph(); para_spacing(spacer, 6, 6)

        illus = cover.get('illustration', '')
        if illus and os.path.isfile(illus):
            ip = out.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                cropped = os.path.join(tmp_dir, 'cover_illustration.png')
                cfn, cw, ch, ow, oh = autocrop_image_file(illus, cropped)
                ip.add_run().add_picture(cfn, width=Cm(11))
            except Exception:
                pass
            para_spacing(ip, 6, 12)

        rule = out.add_paragraph()
        add_bottom_border(rule, color=PAL['primary'], sz=18)
        para_spacing(rule, 4, 4)

        tp = out.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = tp.add_run(cover.get('title', ''))
        set_font(r, FONT, size=26, bold=True, color=C_PRIMARY)
        para_spacing(tp, 10, 4)

        if cover.get('subtitle'):
            sp = out.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = sp.add_run(cover['subtitle'])
            set_font(r, FONT, size=13, bold=True, color=C_DARK)
            para_spacing(sp, 4, 20)

        for field in cover.get('fields', []):
            fp2 = out.add_paragraph()
            r1 = fp2.add_run(f"{field}: ")
            set_font(r1, FONT, size=12, bold=False, color=C_DARK)
            r2 = fp2.add_run('_' * 55)
            set_font(r2, FONT, size=12, color=C_DARK)
            para_spacing(fp2, 10, 10)

        insert_page_break(out)

    def build_front_matter():
        for section in front_matter:
            heading = section.get('heading', '')
            if heading:
                hp2 = out.add_paragraph()
                style_heading(hp2, heading)
            for para_text in section.get('paragraphs', []):
                bp = out.add_paragraph()
                style_body(bp, para_text)
        if front_matter:
            insert_page_break(out)

    def build_contents_page():
        if not add_contents:
            return
        tp = out.add_paragraph()
        style_title(tp, 'Contents', page_break=False)
        cp = out.add_paragraph()
        add_toc_field(cp)
        insert_page_break(out)

    build_cover_page()
    build_front_matter()
    build_contents_page()

    # ---- walk source document body ----
    last_heading_type = None
    image_counter = 0
    callout_active = False
    callout_items = []
    stats = {'title': 0, 'heading': 0, 'subheading': 0, 'objective': 0,
              'dialogue_para': 0, 'body': 0, 'images': 0, 'callout': 0,
              'table_vocab': 0, 'table_dialogue': 0, 'table_wordgrid': 0,
              'table_blank': 0, 'table_generic': 0}

    def flush_callout():
        nonlocal callout_active, callout_items
        if callout_items:
            insert_callout_box(callout_items)
            stats['callout'] += 1
        callout_active = False
        callout_items = []

    body = src_doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            para = Paragraph(child, src_doc)
            text = para.text.strip()

            if not text:
                blob, w_emu, h_emu = get_image_blob(src_doc, child)
                if blob:
                    image_counter += 1
                    stats['images'] += 1
                    fn = os.path.join(tmp_dir, f'img_{image_counter}_orig.png')
                    with open(fn, 'wb') as f:
                        f.write(blob)
                    cropped_fn = os.path.join(tmp_dir, f'img_{image_counter}.png')
                    cfn, cw, ch, ow, oh = autocrop_image_file(fn, cropped_fn)
                    w_emu, h_emu = cropped_extent(w_emu, h_emu, ow, oh, cw, ch)
                    p = out.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        run = p.add_run()
                        max_w = Cm(6.0)
                        if w_emu and h_emu:
                            if Emu(w_emu) > max_w:
                                ratio = h_emu / w_emu
                                run.add_picture(cfn, width=max_w, height=Emu(int(max_w * ratio)))
                            else:
                                run.add_picture(cfn, width=Emu(w_emu), height=Emu(h_emu))
                        else:
                            run.add_picture(cfn, width=Cm(1.2))
                    except Exception:
                        pass
                continue

            src_bold_para = any(r.bold for r in para.runs if r.text.strip())
            kind = classify_para(text, src_bold=src_bold_para)
            core_lower = strip_lead_numbering(text).lower()

            if kind in ('title', 'heading'):
                flush_callout()

            if kind == 'title':
                p = out.add_paragraph()
                style_title(p, text)
                last_heading_type = 'title'
                stats['title'] += 1

            elif kind == 'objective_inline':
                p = out.add_paragraph()
                m = re.match(r'^((?:Objective|Learning Objective)\s*:)\s*(.*)$', text, re.IGNORECASE)
                style_objective(p, m.group(1), m.group(2))
                last_heading_type = None
                stats['objective'] += 1

            elif kind == 'heading':
                p = out.add_paragraph()
                style_heading(p, text)
                last_heading_type = 'heading'
                stats['heading'] += 1
                if any(core_lower.startswith(kw) for kw in CALLOUT_TRIGGERS):
                    callout_active = True
                    callout_items = []

            elif kind == 'subheading':
                if callout_active:
                    callout_items.append((text, True, False))
                else:
                    p = out.add_paragraph()
                    style_subheading(p, text)
                stats['subheading'] += 1

            elif kind == 'dialogue_para':
                if callout_active:
                    callout_items.append((text, False, True))
                else:
                    p = out.add_paragraph()
                    style_dialogue_para(p, text)
                stats['dialogue_para'] += 1

            else:
                if callout_active:
                    callout_items.append((text, src_bold_para, False))
                    stats['body'] += 1
                else:
                    p = out.add_paragraph()
                    style_body(p, text, bold=src_bold_para)
                    stats['body'] += 1

        elif tag == 'tbl':
            flush_callout()
            t = Table(child, src_doc)
            kind = classify_table(t)
            nrows, ncols = len(t.rows), len(t.columns)

            if kind == 'blank_lines':
                stats['table_blank'] += 1
                newt = out.add_table(rows=nrows, cols=ncols)
                newt.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(newt, color=border_color, sz=border_sz)
                for r in newt.rows:
                    r.height = Pt(26)
                    set_row_cant_split(r)
                    for c in r.cells:
                        set_cell_margins(c, 60, 60, 120, 120)

            elif kind == 'wordgrid':
                stats['table_wordgrid'] += 1
                newt = out.add_table(rows=nrows, cols=ncols)
                newt.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(newt, color=border_color, sz=border_sz)
                set_col_widths(newt, [1.0 / ncols] * ncols, total_cm=14.0)
                for ri, row in enumerate(t.rows):
                    set_row_cant_split(newt.rows[ri])
                    for ci, cell in enumerate(row.cells):
                        ncell = newt.rows[ri].cells[ci]
                        ncell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        p = ncell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(cell.text.strip())
                        set_font(r, FONT, size=11, bold=True, color=C_DARK)
                        para_spacing(p, 2, 2)
                        set_cell_margins(ncell, 40, 40, 40, 40)

            elif kind in ('vocab', 'dialogue', 'generic'):
                stats[f'table_{kind}'] += 1
                newt = out.add_table(rows=nrows, cols=ncols)
                newt.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(newt, color=border_color, sz=border_sz)

                if ncols == 2:
                    ratios = vocab_cfg.get('column_width_ratios', [0.45, 0.55])
                elif ncols == 3:
                    ratios = grammar_cfg.get('column_width_ratios', [0.3, 0.3, 0.4])
                elif ncols == 4:
                    ratios = [0.30, 0.30, 0.30, 0.10]
                else:
                    ratios = [1.0 / ncols] * ncols
                set_col_widths(newt, ratios, total_cm=18.0)

                header_bg = vocab_cfg.get('header_bg', PAL['primary'])
                header_fg = hexcolor(vocab_cfg.get('header_text_color', '#FFFFFF'))
                alt_bg = vocab_cfg.get('alternate_row_bg', '#F8F9FA')

                for ri, row in enumerate(t.rows):
                    is_header = (kind == 'vocab' and ri == 0)
                    zebra = (ri % 2 == 1)
                    set_row_cant_split(newt.rows[ri])
                    if is_header:
                        trPr = newt.rows[ri]._tr.get_or_add_trPr()
                        trPr.append(OxmlElement('w:tblHeader'))
                    for ci, cell in enumerate(row.cells):
                        ncell = newt.rows[ri].cells[ci]
                        ncell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        set_cell_margins(ncell, 100, 100, 140, 140)
                        p = ncell.paragraphs[0]
                        txt = cell.text.strip()
                        cell_runs = [r for cp in cell.paragraphs for r in cp.runs if r.text.strip()]
                        src_bold = bool(cell_runs) and all(r.bold for r in cell_runs)

                        if is_header:
                            shade_cell(ncell, header_bg)
                            r = p.add_run(txt)
                            set_font(r, FONT, size=10.5, bold=True, color=header_fg)
                        else:
                            if kind == 'vocab' and zebra:
                                shade_cell(ncell, alt_bg)
                            elif kind == 'dialogue' and ri % 2 == 0:
                                shade_cell(ncell, PAL.get('light_neutral', '#DCDDDF'))
                            is_non_latin = bool(NON_LATIN_RE.search(txt))
                            m = DIALOGUE_SPLIT_RE.match(txt) if (kind == 'dialogue' and ci == 0) else None
                            if m:
                                r1 = p.add_run(m.group(1) + ' ')
                                set_font(r1, font_for(m.group(1)), size=10.5, bold=True, color=C_PRIMARY)
                                r2 = p.add_run(m.group(3))
                                set_font(r2, font_for(m.group(3)), size=10.5, color=C_DARK)
                            else:
                                r = p.add_run(txt)
                                if is_non_latin:
                                    kcfg = STY.get('kannada_translation', {})
                                    set_font(r, FONT_INDIC, size=kcfg.get('font_size_pt', 9.5),
                                              bold=src_bold, color=hexcolor(kcfg.get('color', PAL['text_muted'])))
                                else:
                                    set_font(r, FONT, size=10.5, bold=src_bold, color=C_DARK)
                        para_spacing(p, 2, 2, line=STY.get('body_text', {}).get('line_height_pt'))
            else:
                stats['table_generic'] += 1
                newt = out.add_table(rows=nrows, cols=ncols)
                set_table_borders(newt, color=border_color, sz=border_sz)
                for ri, row in enumerate(t.rows):
                    for ci, cell in enumerate(row.cells):
                        ncell = newt.rows[ri].cells[ci]
                        p = ncell.paragraphs[0]
                        r = p.add_run(cell.text.strip())
                        set_font(r, FONT, size=10.5)

            spacer = out.add_paragraph()
            para_spacing(spacer, 0, 6)

    flush_callout()

    out.save(output_path)
    return stats


# ============================================================================
# 6. CLI
# ============================================================================

def convert_to_pdf(docx_path):
    out_dir = os.path.dirname(os.path.abspath(docx_path)) or '.'
    try:
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', out_dir, docx_path],
            check=True, capture_output=True, timeout=180,
        )
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        return pdf_path if os.path.exists(pdf_path) else None
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
        print(f"  ! PDF conversion skipped ({e})", file=sys.stderr)
        return None


def main():
    ap = argparse.ArgumentParser(
        description="Universal DOCX DTP / standardization pipeline.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    ap.add_argument('input_docx', nargs='?', default=INPUT_DOCX,
                     help="Path to the raw, unedited input .docx "
                          f"(default: the INPUT_DOCX variable at the top of this file, currently {INPUT_DOCX!r})")
    ap.add_argument('-s', '--style', default=STYLE_JSON,
                     help="Path to a JSON design-system file "
                          "(default: the STYLE_JSON variable at the top of this file; "
                          "built-in defaults are used if that is also None)")
    ap.add_argument('-o', '--output', default=OUTPUT_DOCX,
                     help="Path for the formatted output .docx "
                          "(default: the OUTPUT_DOCX variable at the top of this file, or "
                          "<input>_formatted.docx if that is None)")
    ap.add_argument('--pdf', action='store_true', default=MAKE_PDF,
                     help="Also render a PDF of the output via LibreOffice, if available "
                          "(default: the MAKE_PDF variable at the top of this file)")
    args = ap.parse_args()

    if not os.path.isfile(args.input_docx):
        ap.error(f"input file not found: {args.input_docx}\n"
                  f"(edit INPUT_DOCX near the top of dtp_pipeline.py, or pass a path as an argument)")

    cfg = DEFAULT_CONFIG
    if args.style:
        if not os.path.isfile(args.style):
            ap.error(f"style JSON not found: {args.style}")
        with open(args.style, encoding='utf-8') as f:
            user_cfg = json.load(f)
        cfg = deep_merge(DEFAULT_CONFIG, user_cfg)

    output_path = args.output
    if not output_path:
        root, _ = os.path.splitext(args.input_docx)
        output_path = f"{root}_formatted.docx"

    print(f"Input:  {args.input_docx}")
    print(f"Style:  {args.style or '(built-in defaults)'}")
    print(f"Output: {output_path}")
    print(f"Cover page:     {'on' if COVER.get('enabled') else 'off'}")
    print(f"Front matter:   {len(FRONT_MATTER)} section(s)")
    print(f"Contents page:  {'on' if ADD_CONTENTS_PAGE else 'off'}")
    print("Processing...")

    with tempfile.TemporaryDirectory() as tmp_dir:
        stats = build(args.input_docx, output_path, cfg, tmp_dir,
                       cover=COVER, front_matter=FRONT_MATTER, add_contents=ADD_CONTENTS_PAGE)

    print("Done.")
    print("  Titles:            ", stats['title'])
    print("  Section headings:  ", stats['heading'])
    print("  Sub-headings:      ", stats['subheading'])
    print("  Objective/callout boxes:", stats['objective'] + stats['callout'])
    print("  Dialogue lines:    ", stats['dialogue_para'])
    print("  Body paragraphs:   ", stats['body'])
    print("  Images placed:     ", stats['images'])
    print("  Vocab tables:      ", stats['table_vocab'])
    print("  Dialogue tables:   ", stats['table_dialogue'])
    print("  Word-search grids: ", stats['table_wordgrid'])
    print("  Blank-line tables: ", stats['table_blank'])
    print("  Generic tables:    ", stats['table_generic'])

    if args.pdf:
        print("Converting to PDF...")
        pdf_path = convert_to_pdf(output_path)
        if pdf_path:
            print(f"PDF:    {pdf_path}")

    print(f"\nSaved: {output_path}")


if __name__ == '__main__':
    main()
